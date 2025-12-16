import google.generativeai as genai
from PIL import Image
import io
from config import Config 

TEXT_PRIORITY = [
    'gemini-2.5-flash',
    'gemini-2.5-pro',
    'gemini-flash-latest',
    'gemini-pro-latest',
    'gemini-2.5-flash-lite',
    'gemini-2.0-flash',
    'gemma-3-12b-it',
    'text-bison-001'
]

class GeminiAI:
    def __init__(self):
        genai.configure(api_key=Config.GEMINI_API_KEY)
        self.candidates = []
        self.model = None
        try:
            try:
                preferred = getattr(Config, 'GEMINI_MODEL', None)
                if preferred:
                    self.model = genai.GenerativeModel(preferred)
                    return
            except Exception:
                self.model = None

            models = genai.list_models()
            norm = []
            for m in models:
                if isinstance(m, dict):
                    name = m.get('name')
                    methods = m.get('supportedGenerationMethods', [])
                else:
                    name = getattr(m, 'name', None)
                    methods = getattr(m, 'supportedGenerationMethods', []) or []
                if name and any(x in methods for x in ('generateContent', 'generate_content')):
                    short = name.split('/')[-1]
                    norm.append(short)

            norm_text = [n for n in norm if 'image' not in n and 'vision' not in n]

            priority = []
            cfg_model = getattr(Config, 'GEMINI_MODEL', None)
            if cfg_model:
                priority.append(cfg_model.split('/')[-1])
            priority.extend(TEXT_PRIORITY)

            chosen = None
            for p in priority:
                if p in norm_text:
                    chosen = p
                    break
            if not chosen and norm_text:
                chosen = norm_text[0]
            if not chosen and norm:
                chosen = norm[0]

            if chosen:
                try:
                    self.model = genai.GenerativeModel(chosen)
                except Exception:
                    self.model = None
        except Exception:
            self.model = None

    def _try_generate_with_model(self, model, prompt):
        if hasattr(model, 'generate_content'):
            return model.generate_content(prompt)
        if hasattr(model, 'generate_text'):
            return model.generate_text(prompt)
        try:
            return genai.generate_text(model=model.name if hasattr(model, 'name') else model, prompt=prompt)
        except Exception:
            raise
    
    def _summarize_exc(self, e):
        try:
            s = str(e)
        except Exception:
            s = repr(e)
        if 'bytearray' in s:
            return type(e).__name__
        if len(s) > 300:
            return s[:300] + '...'
        return s

    def _no_format_prefix(self) -> str:
        return (
            "На мой запрос отправь ответ без никакого форматирования (кроме выделения кусков кода). "
            "Под форматированием я имею ввиду выделение жирным, курсивом и тд. А абзацы и прочее можешь добавлять. Также будь вежлив. Запрос: "
        )

    def _ensure_text_model(self):
        try:
            name = getattr(self.model, 'model_name', None) or getattr(self.model, 'name', '')
            short = str(name).split('/')[-1]
        except Exception:
            short = ''
        if short and ('image' in short or 'vision' in short):
            for p in TEXT_PRIORITY:
                try:
                    self.model = genai.GenerativeModel(p)
                    return
                except Exception:
                    continue

    def chat(self, message, history=None):
        try:
            if not self.model:
                return (
                    "Ошибка ИИ: не найдена доступная модель для генерации. "
                    "Проверьте правильность `GEMINI_API_KEY` и доступность моделей в вашем аккаунте."
                )
            self._ensure_text_model()
            no_format_prefix = self._no_format_prefix()
            if history:
                prompt_lines = []
                for m in history[-10:]:
                    role = m.get('role', 'user')
                    content = m.get('content', '')
                    if role == 'user':
                        prompt_lines.append(f"User: {content}")
                    else:
                        prompt_lines.append(f"Assistant: {content}")
                prompt_lines.append(f"User: {message}\nAssistant:")
                prompt = no_format_prefix + "\n" + "\n".join(prompt_lines)
                try:
                    response = self._try_generate_with_model(self.model, prompt)
                except Exception as e:
                    last_exc = e
                    response = None
                    err_text = str(e)
                    candidates_to_try = list(self.candidates)
                    if 'quota' in err_text.lower() or '429' in err_text:
                        fallback_models = [
                            'models/gemini-flash-lite-latest',
                            'models/gemini-2.5-flash-lite',
                            'models/gemma-3-1b-it',
                        ]
                        candidates_to_try = fallback_models + candidates_to_try
                    for mname in candidates_to_try:
                        if any(x in mname for x in ('image', 'vision')):
                            continue
                        try:
                            candidate = genai.GenerativeModel(mname)
                            response = self._try_generate_with_model(candidate, prompt)
                            self.model = candidate
                            break
                        except Exception as e2:
                            last_exc = e2
                            continue
                    if response is None:
                        raise last_exc
            else:
                try:
                    response = self._try_generate_with_model(self.model, no_format_prefix + str(message))
                except Exception as e:
                    last_exc = e
                    response = None
                    err_text = str(e)
                    candidates_to_try = list(self.candidates)
                    if 'quota' in err_text.lower() or '429' in err_text:
                        fallback_models = [
                            'models/gemini-2.5-flash-lite',
                            'models/gemini-flash-lite-latest',
                            'models/gemini-2.5-flash-lite',
                            'models/gemma-3-1b-it',
                        ]
                        candidates_to_try = fallback_models + candidates_to_try
                    for mname in candidates_to_try:
                        if any(x in mname for x in ('image', 'vision')):
                            continue
                        try:
                            candidate = genai.GenerativeModel(mname)
                            response = self._try_generate_with_model(candidate, no_format_prefix + str(message))
                            self.model = candidate
                            break
                        except Exception as e2:
                            last_exc = e2
                            continue
                    if response is None:
                        raise last_exc
            return response.text
        except Exception as e:
            return f"Ошибка ИИ: {str(e)}"
    
    def analyze_image(self, image_bytes, prompt=None):
        try:
            image = Image.open(io.BytesIO(image_bytes))
            no_format_prefix = self._no_format_prefix()
            if not prompt:
                prompt = (
                    "Проанализируй это изображение. Если это рукописный текст - перепиши его в печатном виде. "
                    "Если это математический пример или задача - реши её и объясни решение. "
                    "Если это что-то еще - опиши что видишь и дай рекомендации."
                )
            prompt = no_format_prefix + prompt
            try:
                response = self.model.generate_content([prompt, image])
                return response.text
            except Exception as e:
                err_text = str(e)
                if 'quota' in err_text.lower() or 'quota exceeded' in err_text.lower() or '429' in err_text:
                    fallback_models = [
                        'models/gemini-2.5-flash-lite',
                        'models/gemini-flash-lite-latest',
                        'models/gemini-2.5-flash-image-preview',
                        'models/nano-banana-pro-preview',
                    ]
                    last_exc = e
                    for fm in fallback_models:
                        try:
                            candidate = genai.GenerativeModel(fm)
                            resp = candidate.generate_content([prompt, image])
                            self.model = candidate
                            return resp.text
                        except Exception as e2:
                            last_exc = e2
                            continue
                    return f"Видимо квота превышена! Ошибка анализа изображения: {err_text} (запасные модели не помогли: {str(last_exc)})"
                return f"Видимо квота превышена! Ошибка анализа изображения: {err_text}"
        except Exception as e:
            return f"Ошибка анализа изображения: {str(e)}"
    
    def analyze_pdf(self, pdf_bytes):
        try:
            no_format_prefix = self._no_format_prefix()
            prompt = no_format_prefix + (
                "Проанализируй этот PDF документ. Если это учебный материал, сделай краткий конспект."
            )
            response = self.model.generate_content([prompt, pdf_bytes])
            return response.text
        except Exception as e:
            return f"Ошибка анализа PDF: {str(e)}"
    
    def analyze_document(self, file_bytes, file_name, prompt=None):
        import tempfile
        
        try:
            text_payload = None
            try:
                if file_name.lower().endswith('.txt') and isinstance(file_bytes, (bytes, bytearray)):
                    try:
                        text_payload = file_bytes.decode('utf-8')
                    except Exception:
                        try:
                            text_payload = file_bytes.decode('cp1251')
                        except Exception:
                            text_payload = None
                elif file_name.lower().endswith('.pdf'):
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(io.BytesIO(file_bytes))
                        pages = []
                        for p in reader.pages:
                            try:
                                pages.append(p.extract_text() or '')
                            except Exception:
                                pages.append('')
                        text_payload = "\n".join(pages).strip()
                        if not text_payload:
                            text_payload = None
                    except Exception:
                        text_payload = None
                elif file_name.lower().endswith('.docx'):
                    try:
                        from docx import Document
                        from io import BytesIO
                        doc = Document(BytesIO(file_bytes))
                        paras = [p.text for p in doc.paragraphs]
                        text_payload = "\n".join(paras).strip()
                        if not text_payload:
                            text_payload = None
                    except Exception:
                        text_payload = None
            except Exception:
                text_payload = None

            no_format_prefix = self._no_format_prefix()

            if text_payload:
                try:
                    max_len = 150000
                    if len(text_payload) > max_len:
                        text_payload = text_payload[:max_len]
                    response = self.model.generate_content([(no_format_prefix + (prompt or '')), text_payload])
                    return response.text
                except Exception as e:
                    err_text = self._summarize_exc(e)

            with tempfile.NamedTemporaryFile(suffix=file_name[file_name.rfind('.'):] if '.' in file_name else '', delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            if not prompt:
                prompt = (
                    "Проанализируй этот документ. "
                    "Если это текст - кратко суммаризируй основные моменты. "
                    "Если это учебный материал - сделай краткий конспект. "
                    "Если это данные/таблица - объясни что ты видишь. "
                    "Будь краток и информативен."
                )
            prompt = no_format_prefix + prompt
            
            try:
                uploaded_file = None
                if hasattr(genai, 'upload_file'):
                    try:
                        uploaded_file = genai.upload_file(tmp_path)
                        response = self.model.generate_content([prompt, uploaded_file])
                        return response.text
                    except Exception as e:
                        err_text = str(e)
                else:
                    err_text = ''
                    text_payload = None
                    try:
                        if isinstance(file_bytes, (bytes, bytearray)):
                            try:
                                text_payload = file_bytes.decode('utf-8')
                            except Exception:
                                try:
                                    text_payload = file_bytes.decode('cp1251')
                                except Exception:
                                    text_payload = None
                    except Exception:
                        text_payload = None

                    if text_payload is None and file_name.lower().endswith('.pdf'):
                        try:
                            from PyPDF2 import PdfReader
                            reader = PdfReader(io.BytesIO(file_bytes))
                            pages = []
                            for p in reader.pages:
                                try:
                                    pages.append(p.extract_text() or '')
                                except Exception:
                                    pages.append('')
                            text_payload = "\n".join(pages).strip()
                        except Exception:
                            text_payload = None

                    if text_payload is None and file_name.lower().endswith('.docx'):
                        try:
                            from docx import Document
                            from io import BytesIO
                            doc = Document(BytesIO(file_bytes))
                            paras = [p.text for p in doc.paragraphs]
                            text_payload = "\n".join(paras).strip()
                        except Exception:
                            text_payload = None

                    if text_payload:
                        try:
                            max_len = 150000
                            if len(text_payload) > max_len:
                                text_payload = text_payload[:max_len]
                            response = self.model.generate_content([prompt, text_payload])
                            return response.text
                        except Exception as e:
                            err_text = str(e)
                    else:
                        err_text = (
                            "Клиент Gemini не поддерживает загрузку файлов (upload_file), "
                            "и из присланного файла не удалось извлечь текст автоматически. "
                            "Пожалуйста, пришлите текстовый файл (.txt), PDF с распознаваемым текстом, "
                            "или скопируйте содержимое в сообщение."
                        )

                if 'quota' in err_text.lower() or 'quota exceeded' in err_text.lower() or '429' in err_text:
                    fallback_models = [
                        'models/gemini-2.5-flash-lite',
                        'models/gemini-flash-lite-latest',
                        'models/gemini-2.5-flash-image-preview',
                        'models/gemma-3-1b-it',
                        'models/nano-banana-pro-preview'
                    ]
                    uploaded_file_fallback = None
                    if hasattr(genai, 'upload_file'):
                        try:
                            uploaded_file_fallback = genai.upload_file(tmp_path)
                        except Exception:
                            uploaded_file_fallback = None

                    last_exc = None
                    for fm in fallback_models:
                        try:
                            candidate = genai.GenerativeModel(fm)
                            if uploaded_file_fallback is not None:
                                resp = candidate.generate_content([prompt, uploaded_file_fallback])
                            else:
                                if 'text_payload' in locals() and text_payload:
                                    resp = candidate.generate_content([prompt, text_payload])
                                else:
                                    raise Exception('Невозможно передать бинарный файл напряму��; извлеките текст или используйте upload_file')
                            self.model = candidate
                            return resp.text
                        except Exception as e2:
                            last_exc = e2
                            continue
                    return f"Ошибка анализа документа: {err_text} (запасные модели не помогли: {self._summarize_exc(last_exc)})"
                return f"Ошибка анализа документа: {self._summarize_exc(err_text)}"
            except Exception as e:
                return f"Ошибка анализа документа: {self._summarize_exc(e)}"
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except Exception:
                pass
